from typing import List, Dict, Any, Optional, Tuple
import os
from pathlib import Path
from datetime import datetime
import hashlib

from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain.docstore.document import Document as LangchainDocument

class DocumentProcessor:
    """Handles document processing operations"""
    
    def __init__(self, config: dict):
        self.config = config
        self.chunk_size = config['documents']['chunk_size']
        self.chunk_overlap = config['documents']['chunk_overlap']
        self.max_file_size = config['documents']['max_file_size'] * 1024 * 1024  # Convert MB to bytes
        self.supported_types = config['documents']['supported_types']
        self.max_tokens = config['documents']['max_tokens']
        
        # Initialize text splitter with optimized parameters
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
    def _get_file_hash(self, file_content: bytes) -> str:
        """Generate unique hash for file content"""
        return hashlib.md5(file_content).hexdigest()
    
    def validate_file(self, file: Any) -> None:
        """Validate uploaded file"""
        if not file:
            raise ValueError("No file provided")
        
        # Check file size
        if hasattr(file, 'size') and file.size > self.max_file_size:
            raise ValueError(f"File size exceeds limit ({self.config['documents']['max_file_size']}MB)")
        
        # Check file type
        file_ext = os.path.splitext(file.name)[1][1:].lower()
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def _extract_pdf_text(self, file: Any) -> str:
        """Extract text from PDF document"""
        try:
            pdf_reader = PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")
    
    def _extract_docx_text(self, file: Any) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error processing Word document: {str(e)}")
    
    def process_document(self, file: Any) -> Tuple[List[LangchainDocument], Dict[str, Any]]:
        """Process a single document and return chunks with metadata"""
        self.validate_file(file)
        
        # Generate file hash
        file_content = file.getvalue()
        file_hash = self._get_file_hash(file_content)
        
        # Extract text based on file type
        file_ext = os.path.splitext(file.name)[1][1:].lower()
        if file_ext == 'pdf':
            text = self._extract_pdf_text(file)
        elif file_ext in ['docx', 'doc']:
            text = self._extract_docx_text(file)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Create metadata
        metadata = {
            'source': file.name,
            'file_hash': file_hash,
            'type': file_ext,
            'processed_at': datetime.now().isoformat(),
            'size': len(file_content),
            'chunks': 0
        }
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        metadata['chunks'] = len(chunks)
        
        # Create Langchain documents with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    **metadata,
                    'chunk': i,
                    'total_chunks': len(chunks)
                }
            )
            documents.append(doc)
        
        return documents, metadata
    
    def process_documents(self, files: List[Any]) -> Tuple[List[LangchainDocument], List[Dict[str, Any]]]:
        """Process multiple documents and return chunks with metadata"""
        if not files:
            raise ValueError("No documents provided")
        
        all_documents = []
        all_metadata = []
        
        for file in files:
            try:
                documents, metadata = self.process_document(file)
                all_documents.extend(documents)
                all_metadata.append(metadata)
            except Exception as e:
                raise ValueError(f"Error processing {file.name}: {str(e)}")
        
        return all_documents, all_metadata
    
    def get_embedding_model(self) -> OpenAIEmbeddings:
        """Get configured embedding model"""
        embedding_config = self.config['documents'].get('embedding_model', 'text-embedding-3-small')
        return OpenAIEmbeddings(model=embedding_config)
    
    def estimate_tokens(self, text: str) -> int:
        """Estimate token count for text"""
        try:
            import tiktoken
            encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
            return len(encoding.encode(text))
        except:
            # Fallback estimation if tiktoken is not available
            return len(text) // 4
    
    def validate_token_limits(self, documents: List[LangchainDocument]) -> List[LangchainDocument]:
        """Validate and filter documents based on token limits"""
        validated_docs = []
        for doc in documents:
            token_count = self.estimate_tokens(doc.page_content)
            if token_count <= self.max_tokens:
                validated_docs.append(doc)
            else:
                # Split large chunk further if possible
                try:
                    sub_chunks = self.text_splitter.split_text(doc.page_content)
                    for chunk in sub_chunks:
                        if self.estimate_tokens(chunk) <= self.max_tokens:
                            new_doc = LangchainDocument(
                                page_content=chunk,
                                metadata={**doc.metadata}
                            )
                            validated_docs.append(new_doc)
                except:
                    # Skip if can't split further
                    continue
        
        return validated_docs